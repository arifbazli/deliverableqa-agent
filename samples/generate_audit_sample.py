from pathlib import Path

import docx

OUTPUT_PATH = Path(__file__).resolve().parent / "audit_sample.docx"


def build() -> None:
    document = docx.Document()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "This audit assessed Meridian Logistics' vendor payment controls for the "
        "period January-December 2025. We identified three control deficiencies, "
        "one of which is rated high risk."
    )

    document.add_heading("Audit Scope and Objectives", level=1)
    document.add_paragraph(
        "This audit covered the accounts payable approval workflow, vendor "
        "master file maintenance, and payment authorization controls across "
        "the finance shared services center. Out of scope: payroll and "
        "treasury functions."
    )

    document.add_heading("Methodology", level=1)
    document.add_paragraph(
        "We tested a random sample of 60 vendor payments (approximately 8% of "
        "total transaction volume for the period), reviewed vendor master file "
        "change logs for the full year, and interviewed the AP approval team "
        "against the documented SOX control matrix."
    )

    document.add_heading("Findings", level=1)
    document.add_paragraph(
        "Finding 1: Vendor master file changes are not consistently reviewed by "
        "a second approver. Of 25 sampled changes, 9 lacked evidence of "
        "secondary review. Risk rating: high."
    )
    document.add_paragraph(
        "Finding 2: Payment approval thresholds are not enforced in the system "
        "for wire transfers over $50,000. Risk rating: medium."
    )
    document.add_paragraph(
        "Finding 3: there were some issues with how invoices get processed."
    )

    document.add_heading("Recommendations", level=1)
    document.add_paragraph(
        "For Finding 1, implement mandatory second-approver sign-off in the "
        "vendor master file system, with quarterly reconciliation."
    )
    document.add_paragraph(
        "For Finding 2, configure the payment system to hard-block wire "
        "transfers over $50,000 without a second authorization."
    )
    document.add_paragraph(
        "For Finding 3, the process should be improved."
    )

    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    print("Planted errors:")
    print("  - consistency: exec summary says 'one high risk finding', but Finding 1 is high AND")
    print("    the summary count of 3 deficiencies doesn't reconcile with risk ratings given (1 high, 1 medium, 1 unrated)")
    print("  - structure: Management Response section is missing entirely (required by checklist)")
    print("  - structure/language_tone: Finding 3 has no control tested, no evidence basis, and its")
    print("    recommendation ('the process should be improved') is vague and non-actionable")


if __name__ == "__main__":
    build()
