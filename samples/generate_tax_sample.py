from pathlib import Path

import docx

OUTPUT_PATH = Path(__file__).resolve().parent / "tax_sample.docx"


def build() -> None:
    document = docx.Document()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "Based on our review, Northwind Manufacturing's proposed equipment "
        "leaseback structure would generate approximately $2.4M in tax "
        "savings over the five-year lease term, driven primarily by the "
        "acceleration of depreciation deductions."
    )

    document.add_heading("Facts and Assumptions", level=1)
    document.add_paragraph(
        "Northwind intends to sell manufacturing equipment with a book value "
        "of $8M to a leasing entity and lease it back over five years. We "
        "assume the leaseback qualifies as a true lease for tax purposes and "
        "that Northwind's marginal tax rate remains at 21% throughout the "
        "lease term."
    )

    document.add_heading("Basis of Advice", level=1)
    document.add_paragraph(
        "This advice is based on IRC Section 168 (MACRS depreciation) and the "
        "sale-leaseback guidance in Rev. Proc. 2001-28, which sets out the "
        "factors the IRS considers in determining whether a leaseback "
        "qualifies as a true lease rather than a financing arrangement."
    )

    document.add_heading("Analysis", level=1)
    document.add_paragraph(
        "Applying Rev. Proc. 2001-28's factors to Northwind's proposed "
        "structure, the arrangement should qualify as a true lease: the "
        "lease term is less than 80% of the equipment's useful life, "
        "Northwind has no fixed-price purchase option, and the leasing "
        "entity bears meaningful residual value risk. Under this "
        "structure, the projected tax benefit is $1.7M over the lease "
        "term, primarily from the leasing entity's ability to claim "
        "accelerated depreciation that is priced into the lease rate."
    )

    document.add_heading("Conclusion and Recommendation", level=1)
    document.add_paragraph(
        "We recommend Northwind proceed with the sale-leaseback structure "
        "as proposed. This will definitely maximize their tax position."
    )

    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    print("Planted errors:")
    print("  - consistency: Executive Summary states $2.4M in tax savings, but the Analysis")
    print("    section's detailed calculation states $1.7M for the same structure")
    print("  - structure: Limitations and Disclaimers section is missing entirely (required")
    print("    by checklist -- no scope/reliance limitations or jurisdiction caveats stated)")
    print("  - language_tone: 'This will definitely maximize their tax position' is an")
    print("    unsubstantiated superlative/absolute claim inappropriate for tax advice")


if __name__ == "__main__":
    build()
