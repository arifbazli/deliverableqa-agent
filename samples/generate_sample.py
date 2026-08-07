from pathlib import Path

import docx

OUTPUT_PATH = Path(__file__).resolve().parent / "advisory_sample.docx"


def build() -> None:
    document = docx.Document()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "Acme Retail Co. engaged Deloitte to assess supply chain resilience. "
        "Our analysis found that consolidating regional warehouses would "
        "reduce logistics costs by 18% within the first year."
    )

    document.add_heading("Background and Objectives", level=1)
    document.add_paragraph(
        "Acme Retail Co. operates 12 regional distribution centers across "
        "three states. The engagement objective was to evaluate whether "
        "warehouse consolidation could reduce operating costs without "
        "degrading delivery times."
    )

    document.add_heading("Approach and Methodology", level=1)
    document.add_paragraph(
        "We conducted a cost-to-serve analysis using 18 months of shipment "
        "data, interviewed regional operations managers, and benchmarked "
        "against three comparable retail networks."
    )

    document.add_heading("Findings and Analysis", level=1)
    document.add_paragraph(
        "Our detailed cost model, incorporating facility lease costs, labor, "
        "and transportation, shows that consolidating from 12 to 7 regional "
        "warehouses would reduce annual logistics costs by 12%, driven "
        "primarily by reduced facility overhead and improved truck fill rates."
    )
    document.add_paragraph(
        "This significantly improved the overall efficiency of the network."
    )

    document.add_heading("Recommendations", level=1)
    document.add_paragraph(
        "We recommend Acme Retail Co. consolidate to 7 regional warehouses "
        "over an 18-month phased transition, prioritizing closure of the "
        "three lowest-utilization facilities in the first phase."
    )

    document.add_heading("Risks and Mitigations", level=1)
    document.add_paragraph("There are some risks to consider with this plan.")

    document.add_heading("Next Steps", level=1)
    document.add_paragraph(
        "Acme Retail Co. leadership should approve the phased consolidation "
        "plan and confirm the closure sequence with regional operations by "
        "the end of next month."
    )

    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    print("Planted errors:")
    print("  - consistency: exec summary says 18% cost reduction, findings section says 12%")
    print("  - language_tone: 'significantly improved' unsubstantiated claim + passive voice")
    print("  - structure: Risks and Mitigations section has only one generic sentence")


if __name__ == "__main__":
    build()
