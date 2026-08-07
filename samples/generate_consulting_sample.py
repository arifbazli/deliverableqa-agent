from pathlib import Path

import docx

OUTPUT_PATH = Path(__file__).resolve().parent / "consulting_sample.docx"


def build() -> None:
    document = docx.Document()

    document.add_heading("Executive Summary", level=1)
    document.add_paragraph(
        "Brightline Insurance engaged us to redesign its claims intake process. "
        "Our recommended changes would cut average claims processing time from "
        "14 days to 6 days, a 57% reduction."
    )

    document.add_heading("Situation and Objectives", level=1)
    document.add_paragraph(
        "Brightline's claims intake process relies on manual document review "
        "and a legacy workflow system. Customer satisfaction scores have "
        "declined as processing times have grown. The objective of this "
        "engagement was to identify root causes of delay and recommend a "
        "redesigned process."
    )

    document.add_heading("Analysis", level=1)
    document.add_paragraph(
        "We mapped the end-to-end claims process across 200 sampled claims "
        "and identified three bottlenecks: manual document classification "
        "(accounting for 4 of the 14 days), sequential rather than parallel "
        "adjuster review (3 days), and a batch-processed approval queue that "
        "only runs twice daily (adding up to 1 day of pure queue time). "
        "Based on this analysis, an automated intake and parallel-review "
        "redesign would reduce average cycle time to approximately 8 days."
    )

    document.add_heading("Recommendations", level=1)
    document.add_paragraph(
        "We recommend Brightline implement automated document classification "
        "using OCR-based triage, move to parallel adjuster review for claims "
        "under $10,000, and move the approval queue to continuous processing."
    )

    document.add_heading("Implementation Roadmap", level=1)
    document.add_paragraph(
        "Phase 1 (Months 1-2): deploy automated document classification. "
        "Phase 2 (Months 3-4): pilot parallel review with two adjuster teams. "
        "Phase 3 (Months 5-6): roll out continuous approval processing and "
        "retire the legacy batch queue."
    )

    document.save(str(OUTPUT_PATH))
    print(f"Wrote {OUTPUT_PATH}")
    print("Planted errors:")
    print("  - consistency: Executive Summary claims 6-day cycle time (57% reduction from 14),")
    print("    but the Analysis section's own bottleneck math only supports ~8 days")
    print("  - structure: Risks and Mitigations section is missing entirely (required by checklist)")
    print("  - structure: no Appendix with supporting data/financial models, despite Analysis")
    print("    referencing a 200-claim sample that implies underlying data")


if __name__ == "__main__":
    build()
